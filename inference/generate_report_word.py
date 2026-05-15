"""
YOLOv8 Toddler Detection - Word Report Generator
Run: python inference/generate_report_word.py
Requires: pip install python-docx
"""

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import PLOTS_DIR, REPORTS_DIR

REPORTS_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = REPORTS_DIR / "toddler_detection_report.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles helpers ────────────────────────────────────────────────────────────
DARK   = RGBColor(0x1a, 0x1a, 0x2e)
ACCENT = RGBColor(0x16, 0x21, 0x3e)
GRAY   = RGBColor(0x55, 0x55, 0x55)
GREEN  = RGBColor(0x1b, 0x5e, 0x20)
WHITE  = RGBColor(0xff, 0xff, 0xff)


def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = WHITE
        # dark background via paragraph shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "1a1a2e")
        pPr.append(shd)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.left_indent  = Cm(0.3)
    elif level == 2:
        run.font.size = Pt(11)
        run.font.color.rgb = ACCENT
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(3)
    return p


def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY
    return p


def kv(key, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(key + ": ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = DARK
    r2 = p.add_run(value)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRAY


def add_table(headers, rows, col_widths_cm, highlight_last=False, header_color="1a1a2e"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"

    # Header row
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].width = Cm(col_widths_cm[i])
        set_cell_bg(hdr_cells[i], header_color)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    # Data rows
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        is_last = highlight_last and ri == len(rows) - 1
        bg = "d4edda" if is_last else ("f5f5f5" if ri % 2 == 0 else "ffffff")
        for ci, val in enumerate(row):
            cells[ci].width = Cm(col_widths_cm[ci])
            set_cell_bg(cells[ci], bg)
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if is_last:
                run.bold = True
    doc.add_paragraph()


def insert_plot(filename, caption):
    p = PLOTS_DIR / filename
    if p.exists():
        doc.add_picture(str(p), width=Inches(5.5))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in c.runs:
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY
    else:
        p2 = doc.add_paragraph(f"[ {caption} - {filename} not found in {PLOTS_DIR} ]")
        for run in p2.runs:
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xaa, 0x44, 0x44)


# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph("YOLOv8 Toddler Detection")
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].bold = True
t.runs[0].font.size = Pt(28)
t.runs[0].font.color.rgb = DARK

t2 = doc.add_paragraph("Detection  |  Tracking  |  Pose Estimation")
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2.runs[0].font.size = Pt(14)
t2.runs[0].font.color.rgb = GRAY

t3 = doc.add_paragraph("Full Pipeline Project Report")
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
t3.runs[0].font.size = Pt(16)
t3.runs[0].font.color.rgb = ACCENT
t3.runs[0].bold = True

d = doc.add_paragraph(str(date.today()))
d.alignment = WD_ALIGN_PARAGRAPH.CENTER
d.runs[0].font.size = Pt(11)
d.runs[0].font.color.rgb = GRAY

doc.add_paragraph()
abs_p = doc.add_paragraph(
    "This report documents the complete pipeline for fine-tuning YOLOv8s on a custom "
    "toddler detection dataset, covering data preparation, five iterative training runs, "
    "hard negative mining across four categories, multi-object tracking with ByteTrack, "
    "17-keypoint pose estimation, a combined real-time pipeline, and false-positive testing "
    "on 7,500 adult-only video frames."
)
abs_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in abs_p.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
heading("1. Project Overview")
body(
    "The objective is to detect toddlers (young children) in indoor video footage from IP cameras "
    "using a single-class YOLOv8s object detection model. The model outputs bounding boxes around "
    "every toddler visible in each frame and is intended for real-time safety monitoring.\n\n"
    "A major challenge was that early model versions falsely detected adults as toddlers due to the "
    "absence of adult examples during training. This was addressed through an iterative hard negative "
    "mining strategy that progressively introduced adult-containing images with empty label files, "
    "teaching the model to treat adults as background. Four categories of hard negatives were "
    "eventually added: Roboflow adult images, custom adult video frames, and Roboflow datasets for "
    "pets, furniture, indoor rooms, and toys."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. TOOLS & TECHNOLOGIES
# ═══════════════════════════════════════════════════════════════════════════════
heading("2. Tools & Technologies")

heading("A. Detection Model - YOLOv8s", level=2)
body(
    "YOLOv8 (You Only Look Once, version 8) is a single-stage, anchor-free real-time object "
    "detector developed by Ultralytics. Unlike two-stage detectors (e.g. Faster R-CNN), YOLOv8 "
    "processes the full image in a single forward pass, making it suitable for real-time video.\n\n"
    "Architecture breakdown:\n"
    "  Backbone - CSPDarknet with C2f (Cross-Stage Partial with feature fusion) blocks. "
    "Extracts multi-scale feature maps at three resolutions (80x80, 40x40, 20x20 for a 640px input). "
    "The C2f block is a reparameterized bottleneck that improves gradient flow during training.\n"
    "  Neck - PANet (Path Aggregation Network) feature pyramid. Merges backbone feature maps "
    "top-down and bottom-up so the detector has both semantic context (deep layers) and spatial "
    "precision (shallow layers) at every scale.\n"
    "  Head - Decoupled detection head: separate branches for classification (is it a toddler?) "
    "and regression (where exactly is it?). Uses Distribution Focal Loss (DFL) for bounding box "
    "regression, which models the box coordinates as a probability distribution rather than a "
    "single point - this is what drives the high mAP50-95 score (precise box localization).\n"
    "  Anchor-free design: box centres are predicted directly relative to each grid cell with no "
    "predefined anchor shapes, which simplifies training and generalises better to unusual aspect ratios."
)
add_table(
    headers=["YOLOv8 Variant", "Parameters", "GFLOPs", "Reason for choice"],
    rows=[
        ["YOLOv8n (nano)",   "3.2 M",  "8.7",  "Run 1 only - underfitted, mAP50-95=0.310"],
        ["YOLOv8s (small)",  "11.1 M", "28.6", "Runs 2-5 - best accuracy/speed trade-off"],
        ["YOLOv8m (medium)", "25.9 M", "78.9", "Not used - too slow for real-time on CPU"],
    ],
    col_widths_cm=[4.0, 2.5, 2.5, 9.5]
)
kv("Input resolution",   "640 x 640 px (letterbox padded, preserves aspect ratio)")
kv("Output per frame",   "N bounding boxes: (x1, y1, x2, y2, confidence, class_id)")
kv("Confidence thresh.", "0.4 at inference (boxes below this discarded before NMS)")
kv("NMS IoU threshold",  "0.45 (default Ultralytics) - suppresses overlapping boxes")
kv("Final model file",   "models/toddler_detect_s100_best.pt  (21.5 MB)")

heading("B. Pose Estimation Model - YOLOv8n-pose", level=2)
body(
    "YOLOv8n-pose is the nano variant of YOLOv8 extended with a keypoint regression head. "
    "In addition to detecting a bounding box around each person, it regresses 17 (x, y, confidence) "
    "keypoint tuples - one per COCO body joint. The model was pre-trained on the COCO Keypoints "
    "dataset (~250k person instances with annotated skeletons) and requires no fine-tuning for "
    "toddlers because the human skeletal topology is identical across age groups.\n\n"
    "Keypoint head: appended to the same YOLOv8n backbone/neck. For each detected person, it "
    "outputs a 51-dimensional vector (17 joints x 3 values: x, y, visibility confidence). "
    "Keypoints with confidence below 0.4 are not drawn."
)
kv("Model file",       "yolov8n-pose.pt  (~6 MB, auto-downloaded from Ultralytics on first run)")
kv("Keypoints",        "17 COCO joints: nose, eyes, ears, shoulders, elbows, wrists, hips, knees, ankles")
kv("Skeleton edges",   "16 bone connections drawn from the COCO skeleton definition")
kv("Keypoint colors",  "Green=face, Orange=shoulders, Blue gradient=arms, Cyan=legs")
kv("Role in pipeline", "Provides per-joint positions for pose analysis and danger detection")

heading("C. Multi-Object Tracker - ByteTrack", level=2)
body(
    "ByteTrack is a multi-object tracking (MOT) algorithm that assigns each detected object a "
    "persistent integer ID across video frames. It is built directly into the Ultralytics library "
    "and activated with a single parameter (tracker='bytetrack.yaml'), requiring no additional "
    "dependencies. ByteTrack works in two matching passes per frame:\n"
    "  Pass 1 (high-confidence): detections above a high threshold are matched to existing tracks "
    "using the Hungarian algorithm on IoU (Intersection over Union) distance. Matched tracks are "
    "updated with the new box; unmatched tracks enter a lost buffer.\n"
    "  Pass 2 (low-confidence): detections that fall between a floor threshold and the high "
    "threshold are used to recover tracks from the lost buffer. This two-stage design is the "
    "key innovation: it avoids discarding detections that are valid but temporarily uncertain "
    "(occlusion, motion blur), which is the main source of ID switches in other trackers.\n\n"
    "Track lifecycle: New -> Active (matched in Pass 1 or 2) -> Lost (unmatched, kept in buffer "
    "for EXIT_PATIENCE=45 frames) -> Removed (logged as exit event)."
)
kv("Config file",     "bytetrack.yaml (built into Ultralytics installation)")
kv("Matching metric", "IoU distance + Hungarian assignment (linear sum assignment)")
kv("Exit patience",   "45 frames absent before track removed and exit event logged")
kv("Trail rendering", "Last 50 centre positions per track; line opacity proportional to age")
kv("ID color scheme", "Unique HSV hue per ID: hue = (ID x 47) mod 180, converted to BGR")

heading("D. Framework & Runtime Stack", level=2)
add_table(
    headers=["Component", "Version", "Role"],
    rows=[
        ["Python",       "3.10",     "Runtime language"],
        ["Ultralytics",  "8.4.50",   "YOLOv8 training, inference, tracking, pose API"],
        ["PyTorch",      "2.10",     "Deep learning backend (autograd, CUDA tensors)"],
        ["CUDA",         "12.8",     "GPU acceleration (Colab A100 training)"],
        ["OpenCV (cv2)", "4.x",      "Frame decoding, drawing, display, video writing"],
        ["NumPy",        "1.x",      "Array operations for box/keypoint coordinate math"],
        ["fpdf2",        "2.x",      "PDF report generation"],
        ["python-docx",  "1.x",      "Word report generation (this document)"],
        ["Roboflow API", "roboflow", "Downloading hard negative datasets"],
    ],
    col_widths_cm=[3.5, 2.5, 12.5]
)

heading("E. Training Hardware", level=2)
add_table(
    headers=["Run", "Hardware", "Notes"],
    rows=[
        ["1",   "Local CPU (Windows)",       "workers=0, AMP disabled - very slow, 50 epochs"],
        ["2-4", "Google Colab A100 (40 GB)", "~1 hour per 100-epoch run, batch=64, AMP enabled"],
        ["5",   "Google Colab A100 (40 GB)", "Final run - pets/furniture/rooms/toys hard negatives added on Colab"],
    ],
    col_widths_cm=[1.5, 5.0, 12.0]
)

heading("F. YOLO Detection Dataset Format", level=2)
body(
    "Each image has a paired .txt label file with the same stem name. Each line represents one object:\n"
    "  <class_id> <cx> <cy> <w> <h>\n"
    "All coordinates are normalised to [0, 1] relative to the image dimensions. class_id=0 always "
    "(toddler is the only class). Hard negative images have an empty label file (0 bytes) - this "
    "signals to the model that any detection on that image is a false positive that should be "
    "suppressed, teaching the model what NOT to detect."
)
kv("Image format", ".jpg, resized to 640x640 px during training (letterbox, keeps aspect ratio)")
kv("Label format", ".txt - one row per object: class cx cy w h (all values normalised 0-1)")
kv("Hard negatives", "Empty .txt file (0 bytes) - background-only images, no annotations")
kv("Dataset config", "dataset/detect_child/data.yaml - declares train/val/test paths and nc=1")

heading("G. Hard Negative Sources", level=2)
add_table(
    headers=["Source", "Category", "Images", "Purpose"],
    rows=[
        ["Roboflow - Children vs Adults", "Adult humans",         "~300", "Teach model to ignore adult bodies"],
        ["Custom video frames",           "Adults (local env.)",  "74",   "Target specific deployment camera angle"],
        ["Roboflow - Oxford IIIT Pets",   "Cats & dogs",          "~150", "Suppress animal silhouettes (resemble crouching toddlers)"],
        ["Roboflow - Furniture",          "Chairs, tables, sofas","~100", "Suppress common indoor furniture"],
        ["Roboflow - Indoor rooms",       "Walls, floors, windows","~100","Suppress empty room backgrounds"],
        ["Roboflow - Toys",               "Toys on floor",        "~100", "Suppress small objects at floor level"],
    ],
    col_widths_cm=[5.0, 4.0, 2.0, 7.5]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. DATASET PREPARATION
# ═══════════════════════════════════════════════════════════════════════════════
heading("3. Dataset Preparation")
body("The dataset was built in four stages across the five training runs:")
body(
    "Stage 1 - Re-split (70 / 15 / 15)\n"
    "All images and labels were pooled and randomly re-split using seed=42 into 70% train, "
    "15% validation, 15% test. The original naive split had a large val/test gap (0.948 vs 0.700 "
    "mAP50) because validation images came from the same video clips as training images. "
    "After re-splitting the gap dropped to ~2%."
)
body(
    "Stage 2 - Roboflow adult hard negatives\n"
    "Adult-only images from the 'Children vs Adults' Roboflow dataset were downloaded and filtered: "
    "only images with no child annotations were kept. These were added to the train set with empty "
    "label files (background class), teaching the model to suppress adult detections."
)
body(
    "Stage 3 - Custom video hard negatives\n"
    "Frames were extracted (1 per 2 seconds) from real deployment footage containing only adults, "
    "yielding 74 frames after manual review. All were added to the train set with empty labels. "
    "This targeted the specific visual environment and camera angle of the deployment scenario."
)
body(
    "Stage 4 - Additional category hard negatives (Run 5)\n"
    "To further reduce false positives from non-human objects, four additional Roboflow datasets "
    "were downloaded on Google Colab: Oxford IIIT Pets (cats & dogs), a furniture detection "
    "dataset (chairs, tables, sofas), indoor room scenes (walls, floors, windows), and a toys "
    "detection dataset. Person-containing images were filtered out; remaining images were added "
    "to the train set with empty label files."
)
add_table(
    headers=["Split", "Images (Run 4 baseline)", "Hard Negatives (total)", "Labeled Images"],
    rows=[
        ["train", "2018", "723+", "1295"],
        ["val",   "289",  "8",    "281"],
        ["test",  "290",  "7",    "283"],
    ],
    col_widths_cm=[2.5, 5.0, 5.5, 5.5]
)
body(
    "Note: Run 5 added further hard negatives to the train split (pets, furniture, rooms, toys). "
    "Val and test splits remain unchanged from Run 2 onward to keep evaluation consistent."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4. TRAINING CONFIGURATION
# ═══════════════════════════════════════════════════════════════════════════════
heading("4. Training Configuration (Run 5)")
add_table(
    headers=["Parameter", "Value"],
    rows=[
        ["Base weights",      "yolov8s.pt (COCO pretrained - clean start)"],
        ["Epochs",            "100"],
        ["Image size",        "640 x 640 px"],
        ["Batch size",        "64"],
        ["Optimizer",         "AdamW - lr=0.002, momentum=0.9, weight_decay=0.0005"],
        ["Early stopping",    "patience=15 epochs"],
        ["LR schedule",       "Cosine decay (lrf=0.01)"],
        ["Warmup",            "3 epochs"],
        ["Augmentations",     "Mosaic, flip LR (0.5), HSV, Blur, MedianBlur, CLAHE, Erasing"],
        ["AMP",               "Enabled (float16 mixed precision)"],
        ["Close mosaic",      "Last 10 epochs"],
        ["Drive checkpoint",  "Every 10 epochs via on_train_epoch_end callback"],
        ["Training hardware", "Google Colab A100 (40 GB VRAM)"],
        ["Training time",     "~1 hour (A100)"],
    ],
    col_widths_cm=[5.5, 13.0]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 5. TRAINING HISTORY
# ═══════════════════════════════════════════════════════════════════════════════
heading("5. Training History (All 5 Runs)")
add_table(
    headers=["Run", "Description", "Val mAP50", "Test mAP50", "Test mAP50-95"],
    rows=[
        ["1", "YOLOv8n, original naive split",                   "0.948", "0.700", "0.310"],
        ["2", "YOLOv8s, proper 70/15/15 re-split",               "0.956", "0.936", "0.537"],
        ["3", "Run 2 + Roboflow adult hard negatives",            "0.971", "0.951", "0.565"],
        ["4", "Run 3 + custom video hard negatives (74 frames)",  "0.961", "0.951", "0.567"],
        ["5", "Run 4 + pets/furniture/rooms/toys hard negatives", "n/a",   "0.959", "0.575"],
    ],
    col_widths_cm=[1.0, 9.0, 2.5, 2.8, 3.2],
    highlight_last=True
)
body(
    "Key takeaways:\n"
    "  Run 1 -> 2: The val/test gap dropped from 24.8% to 2.0% by fixing the data-split bias.\n"
    "  Run 2 -> 3: Roboflow adult hard negatives raised test mAP50 from 0.936 to 0.951.\n"
    "  Run 3 -> 4: Custom environment-specific hard negatives (74 frames) eliminated adult "
    "false positives while keeping accuracy stable at 0.951.\n"
    "  Run 4 -> 5: Adding pets, furniture, rooms, and toys hard negatives further improved "
    "test mAP50 to 0.959 (+0.8%) and mAP50-95 to 0.575 (+0.8%)."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 6. FINAL MODEL RESULTS
# ═══════════════════════════════════════════════════════════════════════════════
heading("6. Final Model Results (Run 5 - toddler_detect_s100_best.pt)")
body(
    "The saved model is best.pt, automatically selected by Ultralytics as the checkpoint "
    "that achieved the highest validation fitness score across all 100 epochs. Training ran "
    "the full 100 epochs; best.pt corresponds to epoch 89, where validation mAP50 peaked "
    "at 0.9653. After epoch 89 the model continued training but validation metrics plateaued "
    "and slightly fluctuated as the cosine learning rate decayed toward zero. last.pt "
    "(epoch 100) scored marginally lower (mAP50 = 0.9589) and is not used."
)
body("Validation set metrics at best epoch (epoch 89):")
add_table(
    headers=["Metric", "Value", "Meaning"],
    rows=[
        ["Precision",          "94.47%",        "Of all boxes drawn, 94.47% are real toddlers (low FP rate)"],
        ["Recall",             "93.30%",        "Of all real toddlers, 93.30% are detected (low miss rate)"],
        ["F1 Score",           "93.88%",        "Harmonic mean of Precision and Recall"],
        ["mAP50 (val)",        "96.53%",        "Area under P-R curve at IoU >= 0.5 on validation set"],
        ["mAP50-95 (val)",     "56.81%",        "mAP averaged over IoU thresholds 0.50 to 0.95"],
        ["mAP50 (test)",       "95.90%",        "Same metric on the held-out test split (unseen data)"],
        ["mAP50-95 (test)",    "57.50%",        "Box localization precision on unseen data"],
        ["Inference speed",    "~6 ms / image", "On A100 GPU"],
        ["Model file size",    "21.5 MB",       "toddler_detect_s100_best.pt"],
        ["Parameters",         "11,125,971",    "YOLOv8s architecture"],
        ["Confidence thresh.", "0.4",           "Used at inference and FP testing"],
    ],
    col_widths_cm=[3.5, 3.0, 12.0]
)
body(
    "Note on accuracy in object detection: there is no single accuracy metric equivalent "
    "to classification accuracy. Precision, Recall, F1, and mAP are the standard measures. "
    "The F1 score of 93.88% is the closest single-number summary of overall detection quality. "
    "mAP50 of 95.9% on the test set is the primary benchmark metric."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 7. FALSE POSITIVE TESTING
# ═══════════════════════════════════════════════════════════════════════════════
heading("7. False Positive Testing on Adult-Only Videos")
body(
    "After Run 5, the model was tested on 5 adult-only video clips recorded locally "
    "(30-60 seconds each, different lighting and environments). All detections in these "
    "videos are by definition false positives. Inference was run at conf=0.4 using "
    "inference/test_videos.py."
)
add_table(
    headers=["Video", "Frames", "FP Frames", "FP%", "Max Conf", "Status"],
    rows=[
        ["video_2026-05-13_19-29-50.mp4", "1207", "3",  "0.2%", "0.56", "Low FP"],
        ["IMG_3638.mp4",                  "1888", "13", "0.7%", "0.84", "Investigate"],
        ["IMG_1127.mp4",                  "977",  "2",  "0.2%", "0.69", "Low FP"],
        ["IMG_3647.mp4",                  "1606", "0",  "0.0%", "-",    "Clean"],
        ["IMG_3648.MP4",                  "1822", "0",  "0.0%", "-",    "Clean"],
    ],
    col_widths_cm=[5.5, 1.8, 2.0, 1.5, 2.2, 2.5]
)
body(
    "Overall false positive rate: 18 / 7500 frames = 0.24%\n"
    "3 out of 5 videos had zero false positives. IMG_3638.mp4 had 13 FP frames with max "
    "confidence 0.84 - annotated output saved to runs/test_results/IMG_3638/ for review."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 8. PROBLEMS & SOLUTIONS
# ═══════════════════════════════════════════════════════════════════════════════
heading("8. Problems Encountered & Solutions")
add_table(
    headers=["Problem", "Root Cause", "Solution"],
    rows=[
        ["Local train crashed (exit code 255)",
         "Windows PyTorch multiprocessing",
         "workers=0, amp=False, freeze_support()"],
        ["Val/test gap 24.8% (0.948 vs 0.700)",
         "Val images from same video clips as train",
         "Merged all data, re-split 70/15/15 uniformly"],
        ["Adults detected as toddlers (conf 0.84)",
         "No adult examples in training data",
         "Hard negative mining with empty label files"],
        ["Training interrupted at epoch 89",
         "Laptop went to sleep; Colab session kept running",
         "Reconnected to Colab; training had completed on A100"],
        ["Bad resume: 80 COCO classes after interrupt",
         "last.pt optimizer state caused fallback to coco8.yaml",
         "Saved best.pt from Drive, did full clean retrain"],
        ["Drive backup missing on resume",
         "on_train_epoch_end callback not set before resume",
         "Added Drive checkpoint callback to every training cell"],
        ["ImportError: ultralytics callbacks",
         "Unused import from ultralytics.callbacks in cell 7",
         "Removed the unused import line entirely"],
        ["ID fragmentation: 31 IDs in 4 s (before fix)",
         "Detections dropped below conf=0.4 for 1-2 frames causing ByteTrack to lose and reassign tracks",
         "Fixed: DETECT_CONF=0.1, new_track_thresh=0.5, track_buffer=90 in bytetrack_toddler.yaml. "
         "Longest track improved from 121 to 214 frames."],
    ],
    col_widths_cm=[4.5, 5.5, 8.5]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 9. MULTI-OBJECT TRACKING
# ═══════════════════════════════════════════════════════════════════════════════
heading("9. Multi-Object Tracking")
body(
    "After the detection model was validated, ByteTrack was added to assign each toddler a "
    "persistent ID across frames. ByteTrack is built into Ultralytics and requires no extra "
    "dependencies. It is robust to brief occlusions: if two toddlers overlap, their individual "
    "IDs are recovered once they separate.\n\n"
    "ByteTrack Algorithm (two-pass matching per frame):\n"
    "  Pass 1 - High-confidence detections are matched to existing active tracks using IoU-based "
    "Hungarian assignment. Tracks with a strong enough IoU overlap are updated; unmatched tracks "
    "enter a lost state.\n"
    "  Pass 2 - Low-confidence detections (between a floor and the high threshold) are used to "
    "recover lost tracks that were temporarily occluded. This two-pass design is what allows IDs "
    "to survive brief overlaps between two toddlers."
)
kv("Tracker",        "ByteTrack (bytetrack.yaml) - built into Ultralytics")
kv("Script",         "inference/track.py")
kv("Conf threshold", "0.4 (detections below this are rejected entirely)")
kv("ID persistence", "Each toddler keeps its ID across the full video")
kv("Exit patience",  "45 frames of absence before an exit event is logged")
kv("Trail length",   "Last 50 centre positions drawn per track, fading with age")
kv("Trail colors",   "Unique HSV color per track ID - hue = (ID x 47) mod 180")
kv("Output",         "Per-ID coloured box, motion trail, entry/exit log, session summary")
doc.add_paragraph()
body(
    "ByteTrack Tuning - ID Fragmentation Fix (2026-05-15):\n"
    "  Initial test: 31 IDs in 4 s (longest track: 121 frames). Root cause: detections dropped "
    "below conf=0.4 for 1-2 frames, causing ByteTrack to lose and reassign tracks.\n\n"
    "  Fix applied (bytetrack_toddler.yaml):\n"
    "  DETECT_CONF: 0.4 -> 0.1   (passes weak detections to tracker for recovery)\n"
    "  track_high_thresh: 0.25 -> 0.4  (only strong detections drive Pass 1 matching)\n"
    "  track_low_thresh:  0.10 -> 0.05 (weak detections used only for Pass 2 recovery)\n"
    "  new_track_thresh:  0.25 -> 0.50 (requires high confidence to start a new track)\n"
    "  track_buffer:      30   -> 90   (tracks survive 3 s without a detection)\n\n"
    "  After fix: 215 frames | 38 IDs | Longest track: 214 frames (7.1 s)"
)

# ═══════════════════════════════════════════════════════════════════════════════
# 10. POSE ESTIMATION
# ═══════════════════════════════════════════════════════════════════════════════
heading("10. Pose Estimation (Body Part Detection)")
body(
    "YOLOv8n-pose, pre-trained on the COCO dataset, was used to detect 17 body keypoints on "
    "each toddler. No fine-tuning was required - the COCO-pretrained model generalises well to "
    "toddlers since the skeletal structure is the same as for adults."
)
kv("Model",      "yolov8n-pose.pt  (COCO pretrained, ~6 MB, auto-downloaded)")
kv("Keypoints",  "17 COCO joints: nose, eyes, ears, shoulders, elbows, wrists, hips, knees, ankles")
kv("Script",     "inference/pose.py")
kv("Conf thr.",  "0.4 for person detection; 0.4 minimum per-keypoint confidence to draw")
doc.add_paragraph()
add_table(
    headers=["Index", "Keypoint", "Index", "Keypoint"],
    rows=[
        ["0", "Nose",           "9",  "Left wrist"],
        ["1", "Left eye",       "10", "Right wrist"],
        ["2", "Right eye",      "11", "Left hip"],
        ["3", "Left ear",       "12", "Right hip"],
        ["4", "Right ear",      "13", "Left knee"],
        ["5", "Left shoulder",  "14", "Right knee"],
        ["6", "Right shoulder", "15", "Left ankle"],
        ["7", "Left elbow",     "16", "Right ankle"],
        ["8", "Right elbow",    "",   ""],
    ],
    col_widths_cm=[1.5, 5.0, 1.5, 5.0]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 11. COMBINED PIPELINE
# ═══════════════════════════════════════════════════════════════════════════════
heading("11. Combined Detection + Tracking + Pose Pipeline")
body(
    "The three components were fused into a single real-time pipeline. The detection model "
    "acts as a gatekeeper: only regions it confirms as toddlers receive a skeleton from the "
    "pose model. Adults, pets, and background objects are ignored entirely.\n\n"
    "Two models run sequentially on every frame:\n"
    "  1. Detection model + ByteTrack  ->  tracked toddler boxes [(ID, x1, y1, x2, y2), ...]\n"
    "  2. Pose model (full frame)      ->  all visible person skeletons (may include adults)\n"
    "  3. Matching                     ->  skeleton centre must fall inside a tracked toddler box "
    "(+/-30 px margin) to be kept; adult skeletons are silently discarded\n"
    "  4. Drawing                      ->  per-ID colour, motion trail, bounding box + label, "
    "skeleton composited on the same frame\n\n"
    "Matching uses a centre-point test with 30-pixel tolerance:\n"
    "  cx = (pose_x1 + pose_x2) / 2  ;  cy = (pose_y1 + pose_y2) / 2\n"
    "  matched if:  det_x1-30 <= cx <= det_x2+30  AND  det_y1-30 <= cy <= det_y2+30"
)
kv("Script",       "inference/track_and_pose.py")
kv("Detect conf",  "0.4  |  Pose conf: 0.4  |  Keypoint conf threshold: 0.4")
kv("Output video", "runs/track_pose/<stem>_tracked_pose.mp4  (mp4v codec, original resolution)")
kv("Display",      "Live resized window at 960 px width - press Q to quit")
doc.add_paragraph()
body(
    "Live test (2026-05-15) after ByteTrack tuning:\n"
    "  Detection model : toddler_detect_s100_best.pt  |  Tracker: ByteTrack  |  Pose: yolov8n-pose.pt\n"
    "  Total frames    : 215 (7.2 s at 30 fps)\n"
    "  Unique IDs      : 38  |  Longest track: ID 3 (214 frames = 7.1 s)\n"
    "  Output saved to : runs/track_pose/multi_toddler tracking test_tracked_pose.mp4\n\n"
    "All three annotation layers rendered correctly on each confirmed toddler. Adults received "
    "no skeleton, confirming the detection model gatekeeper logic works as intended."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 12. PROJECT STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════════
heading("12. Project Structure")
add_table(
    headers=["File / Directory", "Description"],
    rows=[
        ["config.py",                           "Central path configuration (project root)"],
        ["progress.md",                         "Project status, training history, and next steps"],
        ["notebooks/train_colab.ipynb",         "Google Colab training notebook (9 cells + 6d)"],
        ["training/train_detect.py",            "Local detection training script (CPU / Windows-safe)"],
        ["training/train_pose.py",              "Pose fine-tuning script (for future toddler-specific dataset)"],
        ["inference/detect.py",                 "Detection only - video / image / webcam (conf=0.4)"],
        ["inference/track.py",                  "Detection + ByteTrack tracking with motion trails"],
        ["inference/pose.py",                   "Detection + pose estimation (17 keypoints)"],
        ["inference/detect_and_pose.py",        "Detection + pose, adults filtered out"],
        ["inference/track_and_pose.py",         "Full pipeline: detection + tracking + pose"],
        ["inference/test_videos.py",            "Batch false-positive tester for adult videos"],
        ["inference/generate_report.py",        "PDF report generator (fpdf2)"],
        ["inference/generate_report_word.py",   "Word report generator (python-docx)"],
        ["data/prepare_dataset.py",             "One-time 70/15/15 re-split utility"],
        ["data/extract_hard_negatives.py",      "Extracts frames from video as hard negatives"],
        ["data/download_roboflow_negatives.py", "Downloads pets/furniture/rooms/toys from Roboflow"],
        ["data/download_syrip.py",              "Downloads the SyRIP infant pose dataset"],
        ["data/convert_syrip_to_yolo.py",       "Converts SyRIP COCO keypoints to YOLO pose format"],
        ["models/toddler_detect_s100_best.pt",  "Final detection model weights (Run 5) - 21.5 MB"],
        ["dataset/detect_child/data.yaml",      "YOLO detection dataset configuration"],
        ["adult_negatives/",                    "Custom adult hard negative frames (74 images)"],
        ["plots/",                              "Training plots exported from Google Colab"],
        ["runs/test_results/",                  "Annotated output videos from false-positive testing"],
        ["runs/track_pose/",                    "Annotated output from full pipeline runs"],
    ],
    col_widths_cm=[7.0, 11.5]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 13. NEXT STEPS
# ═══════════════════════════════════════════════════════════════════════════════
heading("13. Next Steps")
add_table(
    headers=["Priority", "Task", "Details"],
    rows=[
        ["1", "Danger detection logic",
         "Use keypoint positions to flag dangerous postures: wrists above nose (climbing), "
         "ankle Y near hip Y (fallen/crawling), large sudden centre-of-mass jump (rapid fall), "
         "box centre near frame edge (leaving monitored zone)."],
        ["3", "Investigate IMG_3638 FPs",
         "Open runs/test_results/IMG_3638/ annotated video, identify objects triggering "
         "conf=0.84 detections, add targeted hard negatives if needed."],
        ["4", "Alert / notification system",
         "Trigger an alert (sound, desktop push) when a danger condition persists for "
         "more than N consecutive frames (debounce to avoid single-frame noise)."],
        ["5", "Pose fine-tuning (optional)",
         "If keypoint accuracy is insufficient, fine-tune yolov8n-pose.pt on a "
         "toddler-specific pose dataset using training/train_pose.py."],
    ],
    col_widths_cm=[1.8, 4.5, 12.2]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 14. TRAINING PLOTS
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("14. Training Plots (Run 5)")
body(
    "All plots were generated by Ultralytics during Run 5 on Google Colab and exported "
    "to the plots/ directory. The vertical dashed line on curve plots marks the best epoch (89)."
)

insert_plot("results.png",                     "Figure 1  - Training & Validation Loss / mAP Curves (100 epochs)")
insert_plot("confusion_matrix_normalized.png", "Figure 2  - Normalized Confusion Matrix (val set)")
insert_plot("confusion_matrix.png",            "Figure 3  - Absolute Confusion Matrix (val set)")
insert_plot("BoxPR_curve.png",                 "Figure 4  - Precision-Recall Curve  |  AP = 0.965")
insert_plot("BoxP_curve.png",                  "Figure 5  - Precision vs Confidence Curve")
insert_plot("BoxR_curve.png",                  "Figure 6  - Recall vs Confidence Curve")
insert_plot("BoxF1_curve.png",                 "Figure 7  - F1 vs Confidence Curve  |  Best F1 = 0.939 @ conf 0.424")
insert_plot("labels.jpg",                      "Figure 8  - Label Distribution & Bounding Box Statistics")
insert_plot("train_batch0.jpg",                "Figure 9  - Training Batch 0: Ground Truth Annotations")
insert_plot("val_batch0_labels.jpg",           "Figure 10 - Validation Batch 0: Ground Truth Labels")
insert_plot("val_batch0_pred.jpg",             "Figure 11 - Validation Batch 0: Model Predictions")
insert_plot("val_batch1_labels.jpg",           "Figure 12 - Validation Batch 1: Ground Truth Labels")
insert_plot("val_batch1_pred.jpg",             "Figure 13 - Validation Batch 1: Model Predictions")
insert_plot("val_batch2_labels.jpg",           "Figure 14 - Validation Batch 2: Ground Truth Labels")
insert_plot("val_batch2_pred.jpg",             "Figure 15 - Validation Batch 2: Model Predictions")

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(str(OUTPUT))
print(f"Report saved to: {OUTPUT}")
